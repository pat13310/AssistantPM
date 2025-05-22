# widgets/DocumentationWidget.py

import sys
import os
import json
import datetime
from pathlib import Path

_current_dir_for_sys_path = os.path.dirname(os.path.abspath(__file__))
_project_root_for_sys_path = os.path.dirname(os.path.dirname(_current_dir_for_sys_path))
if _project_root_for_sys_path not in sys.path:
    sys.path.insert(0, _project_root_for_sys_path)


from PySide6.QtWidgets import (
    QWidget,
    QVBoxLayout,
    QHBoxLayout,
    QTreeView,
    QPushButton,
    QApplication,
    QFileDialog,
    QLineEdit,
    QLabel,
    QProgressBar,
    QMenu,
    QSplitter,
    QComboBox,
    QCheckBox,
    QMessageBox,
    QTabWidget,
    QTextEdit,
    QToolButton,
    QDialog,
    QDialogButtonBox,
    QInputDialog,
)
from PySide6.QtCore import Qt, QThreadPool, Signal, QObject, QTimer, QSize, QModelIndex, QProcess
from PySide6.QtGui import QStandardItemModel, QStandardItem, QIcon, QAction, QColor, QFont
from PySide6.QtWebEngineWidgets import QWebEngineView
from PySide6.QtWebEngineCore import (
    QWebEngineSettings,
    QWebEngineProfile,
    QWebEnginePage,
)

from project.documents.toc import TOC_STRUCTURE
from project.documents.DocType import DocType
from components.widgets.HeaderTitle import HeaderTitle
from agent.OpenAIGenerationTask import OpenAIGenerationTask
from agent.OpenAIStreamingTask import OpenAIStreamingTask
from services.html_renderer import render_html
from services.prompt_builder import build_prompt
from services.export_pdf import export_pdf
from components.dialogues.GitCredentialsDialog import GitCredentialsDialog
from components.dialogues.ProjectNameDlg import ProjectNameDlg
from components.dialogues.PromptEditorDialog import PromptEditorDialog
from components.LazyLoadedComponent import LazyLoadedComponent
import subprocess
import tempfile
import shutil


class GenerationSignals(QObject):
    finished = Signal(str, str)
    error = Signal(str)
    progress = Signal(int)  # Signal pour indiquer la progression (0-100)


class CustomWebPage(QWebEnginePage):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.view_widget = parent

    def createStandardContextMenu(self):
        menu = super().createStandardContextMenu()

        # Supprimer l'action native "View page source"
        for action in menu.actions():
            if action.text().strip().lower() == "view page source":
                menu.removeAction(action)
                break  # Important : une fois supprimée, on sort

        # Ajouter une nouvelle action personnalisée
        if (
            hasattr(self.view_widget, "parent")
            and self.view_widget.parent()
            and hasattr(self.view_widget.parent(), "show_html_source")
        ):
            custom_action = menu.addAction("Voir le code source")
            custom_action.triggered.connect(self.view_widget.parent().show_html_source)
            
            # Ajouter d'autres actions utiles
            if hasattr(self.view_widget.parent(), "edit_content"):
                edit_action = menu.addAction("Éditer le contenu")
                edit_action.triggered.connect(self.view_widget.parent().edit_content)
                
            if hasattr(self.view_widget.parent(), "add_to_favorites"):
                fav_action = menu.addAction("Ajouter aux favoris")
                fav_action.triggered.connect(self.view_widget.parent().add_to_favorites)

        return menu

    def acceptNavigationRequest(self, url, type, isMainFrame):
        # Intercepter les requêtes de navigation vers viewsource://
        if (
            url.scheme() == "viewsource"
            and hasattr(self.view_widget, "parent")
            and self.view_widget.parent()
        ):
            if hasattr(self.view_widget.parent(), "show_html_source"):
                self.view_widget.parent().show_html_source()
                return False
        return super().acceptNavigationRequest(url, type, isMainFrame)
        
    def javaScriptConsoleMessage(self, level, message, lineNumber, sourceID):
        """Surcharge de la méthode pour gérer les messages de la console JavaScript"""
        level_str = "Info"
        if level == QWebEnginePage.WarningMessageLevel:
            level_str = "Warning"
        elif level == QWebEnginePage.ErrorMessageLevel:
            level_str = "Error"
            
        print(f"JavaScript {level_str}: {message} (ligne {lineNumber}, source: {sourceID})")
        
        # Si c'est une erreur et que le widget parent a une méthode pour gérer les erreurs JS
        if level == QWebEnginePage.ErrorMessageLevel and hasattr(self.view_widget, "parent") and self.view_widget.parent():
            parent = self.view_widget.parent()
            if hasattr(parent, "handle_js_error"):
                parent.handle_js_error(message, lineNumber, sourceID)


class DocumentationWidget(QWidget):
    def __init__(self, doc_type: DocType):
        super().__init__()
        self.doc_type = doc_type
        self.thread_pool = QThreadPool()
        self.thread_pool.setMaxThreadCount(2)  # Augmenté pour permettre plusieurs générations
        self.generated_content = {}
        self.current_version = 1
        self.versions = {}  # Stockage des versions {path: {version: content}}
        self.streaming_content = ""  # Stockage temporaire pour le contenu en streaming
        self.is_streaming = False  # Indicateur de génération en streaming
        
        # Créer le dossier de sauvegarde s'il n'existe pas
        self.save_dir = Path(_project_root_for_sys_path) / "saved_docs"
        self.save_dir.mkdir(exist_ok=True)
        
        # Charger les données après avoir créé le dossier de sauvegarde
        self.favorites = self.load_favorites()
        self.history = self.load_history()
        self.custom_prompts = self.load_custom_prompts()
        
        # Initialiser les composants à chargement paresseux
        self._init_lazy_components()

        self.init_ui()
        self.apply_styles()
        
        # CSS par défaut pour le rendu HTML
        self.default_css = """
        body { font-family: 'Segoe UI', Arial, sans-serif; line-height: 1.6; color: #333; }
        h1, h2, h3 { color: #0066cc; }
        pre { background-color: #f5f5f5; padding: 10px; border-radius: 5px; }
        code { font-family: 'Consolas', monospace; }
        .nav-icon { height: 1em; vertical-align: middle; margin-right: 5px; }
        .page-title { margin-bottom: 20px; }
        .page-content { margin-top: 20px; }
        """
        
        # Timer pour auto-sauvegarde
        self.autosave_timer = QTimer(self)
        self.autosave_timer.timeout.connect(self.auto_save)
        self.autosave_timer.start(60000)  # Sauvegarde toutes les minutes
        
    def _init_lazy_components(self):
        """Initialise les composants à chargement paresseux"""
        # Dialogues
        self.prompt_editor_dialog = LazyLoadedComponent(
            lambda: PromptEditorDialog(self)
        )
        self.git_credentials_dialog = LazyLoadedComponent(
            lambda: GitCredentialsDialog(self)
        )
        
        # Vues web
        self.html_view_component = LazyLoadedComponent(
            lambda: self._create_web_view()
        )
        self.version_preview_component = LazyLoadedComponent(
            lambda: self._create_web_view()
        )
    
    def _create_web_view(self):
        """Crée et configure une vue web"""
        view = QWebEngineView()
        
        # Utiliser notre page personnalisée avec menu contextuel étendu
        web_page = CustomWebPage(view)
        view.setPage(web_page)
        
        settings = view.settings()
        
        # Enable JavaScript features
        settings.setAttribute(QWebEngineSettings.JavascriptEnabled, True)
        settings.setAttribute(QWebEngineSettings.LocalContentCanAccessRemoteUrls, True)
        settings.setAttribute(QWebEngineSettings.LocalStorageEnabled, True)
        settings.setAttribute(QWebEngineSettings.WebAttribute.PluginsEnabled, True)
        settings.setAttribute(
            QWebEngineSettings.WebAttribute.AllowRunningInsecureContent, True
        )
        
        # Configure web security
        profile = QWebEngineProfile.defaultProfile()
        profile.setHttpCacheType(QWebEngineProfile.MemoryHttpCache)
        profile.setPersistentCookiesPolicy(QWebEngineProfile.AllowPersistentCookies)
        
        # Initialiser avec un contenu HTML vide mais valide
        view.setHtml("""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Chargement...</title>
            <style>
                body {
                    font-family: 'Segoe UI', sans-serif;
                    display: flex;
                    justify-content: center;
                    align-items: center;
                    height: 100vh;
                    margin: 0;
                    background-color: #f5f5f5;
                }
                .loading {
                    text-align: center;
                    color: #555;
                }
            </style>
        </head>
        <body>
            <div class="loading">
                <h2>Chargement...</h2>
                <p>Le contenu sera affiché ici.</p>
            </div>
        </body>
        </html>
        """)
        
        
        return view
        
        
        
    
    def init_ui(self):
        self.setWindowTitle(f"Documentation IA - {self.doc_type.value}")
        self.setMinimumSize(700, 700)

        self.layout = QVBoxLayout(self)
        self.layout.setSpacing(10)
        self.layout.setContentsMargins(10, 10, 10, 10)

        # Header
        self.setup_header()
        
        # Barre de recherche et outils
        self.setup_toolbar()

        # Body
        self.setup_body()
        
        # Barre de statut
        self.setup_statusbar()

    def setup_header(self):
        icon_map = {
            "CAHIER_DES_CHARGES_FONCTIONNEL": "assets/icons/list-checks.svg",
            "SPECIFICATIONS_FONCTIONNELLES_DETAILLEES": "assets/icons/layers.svg",
            "SPECIFICATIONS_TECHNIQUES_DETAILLEES": "assets/icons/code.svg",
            "STRATEGIE_DE_TESTS_ET_RECETTE": "assets/icons/test-tube.svg",
            "DOSSIER_ARCHITECTURE_TECHNIQUE": "assets/icons/rocket.svg",
        }
        icon_path = icon_map.get(self.doc_type.name, "assets/icons/book-open.svg")
        header = HeaderTitle(title=self.doc_type.value, icon_path=icon_path)
        self.layout.addWidget(header)
    
    def setup_toolbar(self):
        toolbar_layout = QHBoxLayout()
        
        # Barre de recherche
        search_label = QLabel("Rechercher:")
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Rechercher dans la table des matières...")
        self.search_input.textChanged.connect(self.filter_toc)
        self.search_input.setClearButtonEnabled(True)
        
        # Bouton de génération complète
        self.generate_all_btn = QPushButton("Générer tout")
        self.generate_all_btn.setIcon(QIcon("assets/icons/wand.svg"))
        self.generate_all_btn.clicked.connect(self.generate_all_content)
        
        # Bouton de sauvegarde
        self.save_btn = QPushButton("Sauvegarder")
        self.save_btn.setIcon(QIcon("assets/icons/file-down.svg"))
        self.save_btn.clicked.connect(self.save_document)
        
        # Bouton de chargement
        self.load_btn = QPushButton("Charger")
        self.load_btn.setIcon(QIcon("assets/icons/folder-closed.svg"))
        self.load_btn.clicked.connect(self.load_document)
        
        # Bouton de publication Git
        self.git_publish_btn = QPushButton("Publier (Git)")
        self.git_publish_btn.setIcon(QIcon("assets/icons/github.svg"))
        self.git_publish_btn.clicked.connect(self.publish_to_git)
        
        toolbar_layout.addWidget(search_label)
        toolbar_layout.addWidget(self.search_input, 3)
        toolbar_layout.addWidget(self.generate_all_btn, 1)
        toolbar_layout.addWidget(self.save_btn, 1)
        toolbar_layout.addWidget(self.load_btn, 1)
        toolbar_layout.addWidget(self.git_publish_btn, 1)
        
        self.layout.addLayout(toolbar_layout)

    def setup_body(self):
        # Utiliser un QSplitter pour permettre à l'utilisateur d'ajuster les proportions
        self.splitter = QSplitter(Qt.Horizontal)
        
        # Panneau de gauche avec onglets
        self.left_tabs = QTabWidget()
        
        # Onglet Table des matières
        toc_widget = QWidget()
        toc_layout = QVBoxLayout(toc_widget)
        
        # TreeView
        self.toc = QTreeView()
        self.toc.setMouseTracking(True)
        self.toc.setModel(self.build_tree_model())
        self.toc.expandAll()
        self.toc.setHeaderHidden(False)
        self.toc.clicked.connect(self.handle_tree_click)
        self.toc.setContextMenuPolicy(Qt.CustomContextMenu)
        self.toc.customContextMenuRequested.connect(self.show_toc_context_menu)
        
        toc_layout.addWidget(self.toc)
        self.left_tabs.addTab(toc_widget, "Table des matières")
        
        # Onglet Favoris
        favorites_widget = QWidget()
        favorites_layout = QVBoxLayout(favorites_widget)
        self.favorites_list = QTreeView()
        self.favorites_list.setModel(self.build_favorites_model())
        self.favorites_list.clicked.connect(self.handle_favorite_click)
        favorites_layout.addWidget(self.favorites_list)
        self.left_tabs.addTab(favorites_widget, "Favoris")
        
        # Onglet Historique
        history_widget = QWidget()
        history_layout = QVBoxLayout(history_widget)
        self.history_list = QTreeView()
        self.history_list.setModel(self.build_history_model())
        self.history_list.clicked.connect(self.handle_history_click)
        history_layout.addWidget(self.history_list)
        self.left_tabs.addTab(history_widget, "Historique")
        
        # Panneau de droite
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)
        right_layout.setContentsMargins(0, 0, 0, 0)
        
        # Onglets pour le contenu et l'édition
        self.content_tabs = QTabWidget()
        
        # Onglet Visualisation
        view_tab = QWidget()
        view_layout = QVBoxLayout(view_tab)
        view_layout.setContentsMargins(0, 0, 0, 0)
        
        # HTML view with full JavaScript support (lazy loaded)
        self.html_view = self.html_view_component.get()
        view_layout.addWidget(self.html_view)
        self.content_tabs.addTab(view_tab, "Visualisation")
        
        # Onglet Édition
        edit_tab = QWidget()
        edit_layout = QVBoxLayout(edit_tab)
        self.content_editor = QTextEdit()
        self.content_editor.setPlaceholderText("Éditez le contenu HTML ici...")
        edit_layout.addWidget(self.content_editor)
        
        # Boutons d'édition
        edit_buttons = QHBoxLayout()
        self.apply_edits_btn = QPushButton("Appliquer les modifications")
        self.apply_edits_btn.clicked.connect(self.apply_content_edits)
        edit_buttons.addWidget(self.apply_edits_btn)
        
        edit_layout.addLayout(edit_buttons)
        self.content_tabs.addTab(edit_tab, "Édition")
        
        # Onglet Versions
        versions_tab = QWidget()
        versions_layout = QVBoxLayout(versions_tab)
        
        version_header = QHBoxLayout()
        version_label = QLabel("Version:")
        self.version_combo = QComboBox()
        self.version_combo.currentIndexChanged.connect(self.load_version)
        version_header.addWidget(version_label)
        version_header.addWidget(self.version_combo, 1)
        
        versions_layout.addLayout(version_header)
        # Version preview (lazy loaded)
        self.version_preview = self.version_preview_component.get()
        versions_layout.addWidget(self.version_preview)
        
        self.content_tabs.addTab(versions_tab, "Versions")
        
        right_layout.addWidget(self.content_tabs, 1)
        
        # Barre de boutons
        button_layout = QHBoxLayout()
        
        # Bouton de génération
        self.generate_button = QPushButton("Générer")
        self.generate_button.setIcon(QIcon("assets/icons/wand-sparkles.svg"))
        self.generate_button.clicked.connect(self.generate_content)
        
        # Bouton de génération en streaming
        self.generate_streaming_button = QPushButton("Générer (Streaming)")
        self.generate_streaming_button.setIcon(QIcon("assets/icons/zap.svg"))
        self.generate_streaming_button.clicked.connect(self.generate_content_streaming)
        
        # Menu déroulant pour le bouton de génération
        self.generate_menu = QMenu(self)
        self.customize_prompt_action = QAction("Personnaliser le prompt", self)
        self.customize_prompt_action.triggered.connect(self.customize_prompt)
        self.generate_menu.addAction(self.customize_prompt_action)
        
        self.generate_button_menu = QToolButton()
        self.generate_button_menu.setPopupMode(QToolButton.MenuButtonPopup)
        self.generate_button_menu.setMenu(self.generate_menu)
        self.generate_button_menu.setDefaultAction(QAction("Générer", self))
        self.generate_button_menu.setIcon(QIcon("assets/icons/wand-sparkles.svg"))
        self.generate_button_menu.clicked.connect(self.generate_content)
        
        # Bouton d'export
        export_menu = QMenu(self)
        self.export_pdf_action = QAction("Exporter en PDF", self)
        self.export_pdf_action.triggered.connect(self.export_to_pdf)
        export_menu.addAction(self.export_pdf_action)
        
        self.export_word_action = QAction("Exporter en Word", self)
        self.export_word_action.triggered.connect(self.export_to_word)
        export_menu.addAction(self.export_word_action)
        
        self.export_markdown_action = QAction("Exporter en Markdown", self)
        self.export_markdown_action.triggered.connect(self.export_to_markdown)
        export_menu.addAction(self.export_markdown_action)
        
        # Ajouter l'action de publication Git
        self.export_git_action = QAction("Publier sur Git", self)
        self.export_git_action.triggered.connect(self.publish_to_git)
        export_menu.addAction(self.export_git_action)
        
        self.export_button = QToolButton()
        self.export_button.setPopupMode(QToolButton.MenuButtonPopup)
        self.export_button.setMenu(export_menu)
        self.export_button.setDefaultAction(QAction("Exporter", self))
        self.export_button.setIcon(QIcon("assets/icons/file-text.svg"))
        self.export_button.clicked.connect(self.export_to_pdf)
        
        # Bouton pour afficher le code source
        self.show_source_button = QPushButton("Afficher le source")
        self.show_source_button.setIcon(QIcon("assets/icons/code.svg"))
        self.show_source_button.clicked.connect(self.toggle_html_source)
        
        # Bouton pour ajouter aux favoris
        self.favorite_button = QPushButton("Ajouter aux favoris")
        self.favorite_button.setIcon(QIcon("assets/icons/star.svg"))
        self.favorite_button.clicked.connect(self.add_to_favorites)
        
        button_layout.addWidget(self.generate_button_menu)
        button_layout.addWidget(self.generate_streaming_button)
        button_layout.addWidget(self.export_button)
        button_layout.addWidget(self.show_source_button)
        button_layout.addWidget(self.favorite_button)
        
        right_layout.addLayout(button_layout)
        
        # Barre de progression avec style personnalisé
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        self.progress_bar.setVisible(False)
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                border: 1px solid #ddd;
                border-radius: 4px;
                text-align: center;
                height: 20px;
                background-color: #f5f5f5;
            }
            QProgressBar::chunk {
                background-color: #3498db; /* Bleu */
                width: 10px;
                margin: 0.5px;
                border-radius: 2px;
            }
        """)
        right_layout.addWidget(self.progress_bar)
        
        # Ajouter les panneaux au splitter
        self.splitter.addWidget(self.left_tabs)
        self.splitter.addWidget(right_panel)
        self.splitter.setSizes([200, 600])  # Tailles initiales
        
        self.layout.addWidget(self.splitter, 1)
    
    def setup_statusbar(self):
        status_layout = QHBoxLayout()
        
        self.status_label = QLabel("Prêt")
        status_layout.addWidget(self.status_label, 1)
        
        # Indicateur de sauvegarde automatique
        self.autosave_checkbox = QCheckBox("Sauvegarde auto")
        self.autosave_checkbox.setChecked(True)
        self.autosave_checkbox.stateChanged.connect(self.toggle_autosave)
        status_layout.addWidget(self.autosave_checkbox)
        
        # Dernière sauvegarde
        self.last_save_label = QLabel("Dernière sauvegarde: Jamais")
        status_layout.addWidget(self.last_save_label)
        
        self.layout.addLayout(status_layout)

    def apply_styles(self):
        self.setStyleSheet(
            """
            QWidget {
                background-color: #f5f5f5;
                font-family: "Segoe UI", sans-serif;
                color: #333;
            }

            QTreeView {
                background: white;
                border: 1px solid #ddd;
                border-radius: 4px;
                padding: 5px;
                show-decoration-selected: 1; 
            }

            QTreeView::item {
                padding: 4px;
                height: 22px;
            }

            QTreeView::item:hover {
                background: #e8f4fc;
                color: #0078d7;
            }

            QTreeView::item:selected {
                background: #cce8ff;
                color: #0066cc;
            }

            QTreeView::item:selected:hover {
                background: #99d1ff;
            }
            
            QPushButton, QToolButton {
                background-color: #f8f8f8;
                border: 1px solid #ddd;
                border-radius: 4px;
                padding: 6px 12px;
                min-height: 30px;
            }
            
            QPushButton:hover, QToolButton:hover {
                background-color: #e8e8e8;
                border-color: #ccc;
            }
            
            QPushButton:pressed, QToolButton:pressed {
                background-color: #d8d8d8;
            }
            
            QLineEdit {
                padding: 6px;
                border: 1px solid #ddd;
                border-radius: 4px;
                background: white;
            }
            
            QTabWidget::pane {
                border: 1px solid #ddd;
                border-radius: 4px;
                background: white;
            }
            
            QTabBar::tab {
                background: #f0f0f0;
                border: 1px solid #ddd;
                border-bottom-color: #ddd;
                border-top-left-radius: 4px;
                border-top-right-radius: 4px;
                padding: 6px 10px;
                margin-right: 2px;
            }
            
            QTabBar::tab:selected {
                background: white;
                border-bottom-color: white;
            }
            
            QTabBar::tab:hover {
                background: #e0e0e0;
            }
            
            QProgressBar {
                border: 1px solid #ddd;
                border-radius: 4px;
                text-align: center;
                background: white;
            }
            
            QProgressBar::chunk {
                background-color: #4CAF50;
                width: 10px;
                margin: 0.5px;
            }
            
            QTextEdit {
                border: 1px solid #ddd;
                border-radius: 4px;
                background: white;
                font-family: "Consolas", monospace;
            }
            
            QComboBox {
                padding: 5px;
                border: 1px solid #ddd;
                border-radius: 4px;
                background: white;
            }
            
            QCheckBox {
                spacing: 5px;
            }
            
            QCheckBox::indicator {
                width: 16px;
                height: 16px;
            }
            
            QLabel {
                color: #555;
            }
        """
        )

    def build_tree_model(self):
        model = QStandardItemModel()
        model.setHorizontalHeaderLabels(["Table des matières"])
        root = model.invisibleRootItem()

        def add_items(parent, items):
            for title, children in items.items():
                item = QStandardItem(title)
                item.setEditable(False)
                item.setEnabled(True)
                item.setSelectable(True)
                parent.appendRow(item)
                if isinstance(children, dict):
                    add_items(item, children)

        toc = TOC_STRUCTURE.get(self.doc_type.value, {})
        add_items(root, toc)
        return model
    
    def build_favorites_model(self):
        model = QStandardItemModel()
        model.setHorizontalHeaderLabels(["Sections favorites"])
        root = model.invisibleRootItem()
        
        for path in self.favorites:
            item = QStandardItem(path)
            item.setEditable(False)
            item.setData(path)
            root.appendRow(item)
            
        return model
    
    def build_history_model(self):
        model = QStandardItemModel()
        model.setHorizontalHeaderLabels(["Historique des générations"])
        root = model.invisibleRootItem()
        
        for entry in self.history:
            timestamp = entry.get("timestamp", "")
            path = entry.get("path", "")
            display_text = f"{path} ({timestamp})"
            
            item = QStandardItem(display_text)
            item.setEditable(False)
            item.setData(path)
            root.appendRow(item)
            
        return model

    def get_full_path(self, index):
        parts = []
        model = index.model()
        while index.isValid():
            parts.insert(0, model.itemFromIndex(index).text())
            index = index.parent()
        return " > ".join(parts)

    def handle_tree_click(self, index):
        self.current_item_path = self.get_full_path(index)
        print(f"[DEBUG] current_item_path = {self.current_item_path}")
        if self.current_item_path:
            self.load_content(self.current_item_path)
        else:
            print("[WARNING] Aucun chemin sélectionné, chargement annulé.")
        self.add_to_history(self.current_item_path)
    def handle_favorite_click(self, index):
        path = index.data()
        if path:
            self.current_item_path = path
            self.load_content(path)
    
    def handle_history_click(self, index):
        path = index.model().itemFromIndex(index).data()
        if path:
            self.current_item_path = path
            self.load_content(path)
    
    def load_content(self, path):
        html = self.generated_content.get(path)
        if html:
            # Utiliser render_html pour appliquer les styles et icônes
            # Passer skip_title=True pour éviter d'ajouter un titre supplémentaire
            self.html_view.setHtml(render_html(html, self.default_css, skip_title=True))
            self.content_editor.setPlainText(html)
            self.generate_button.setText("Régénérer")
            self.update_version_combo(path)
        else:
            # Remplacer les symboles > par des icônes dans le titre de la page
            formatted_path = path
            if ">" in path:
                formatted_path = path.replace(" > ", " <i class='nav-icon'></i> ")
            
            # Utiliser la même police et style que le reste de l'application
            self.html_view.setHtml(render_html(
                f"<h1 class='page-title'>{formatted_path}</h1><p class='page-content'>Contenu non généré.</p>",
                self.default_css
            ))
            self.content_editor.clear()
            self.generate_button.setText("Générer")
            self.version_combo.clear()
        
        # Mettre à jour le statut avec des icônes au lieu des symboles >
        formatted_status_path = path
        if ">" in path:
            formatted_status_path = path.replace(" > ", " › ")
        self.status_label.setText(f"Section actuelle: {formatted_status_path}")

    def filter_toc(self, text):
        if not text:
            # Si le texte est vide, réinitialiser le modèle
            self.toc.setModel(self.build_tree_model())
            self.toc.expandAll()
            return
        
        # Créer un nouveau modèle pour les résultats filtrés
        filtered_model = QStandardItemModel()
        filtered_model.setHorizontalHeaderLabels(["Résultats de recherche"])
        root = filtered_model.invisibleRootItem()
        # Fonction récursive pour chercher dans le modèle original
        def search_in_model(parent_index, search_text):
            results = []
            model = self.toc.model()
            row_count = model.rowCount(parent_index)
            
            for row in range(row_count):
                index = model.index(row, 0, parent_index)
                text = model.data(index)
                
                # Vérifier si le texte correspond à la recherche
                if search_text.lower() in text.lower():
                    results.append((text, index))
                
                # Rechercher dans les enfants
                child_results = search_in_model(index, search_text)
                if child_results:
                    results.extend(child_results)
            
            return results
        
        # Effectuer la recherche
        results = search_in_model(QModelIndex(), text)
        
        # Ajouter les résultats au modèle filtré
        for text, _ in results:
            item = QStandardItem(text)
            item.setEditable(False)
            root.appendRow(item)
        
        # Appliquer le modèle filtré
        self.toc.setModel(filtered_model)
    
    def generate_content(self):
        if not hasattr(self, "current_item_path"):
            return

        path = self.current_item_path
        
        # Vérifier si un prompt personnalisé existe
        custom_prompt = self.custom_prompts.get(path)
        if custom_prompt:
            prompt = custom_prompt
        else:
            # Extraire le dernier élément du chemin pour le titre de la section
            section_title = path.split("<i class='nav-icon'></i>")[-1].strip() if "<i class='nav-icon'></i>" in path else path
            prompt = build_prompt(self.doc_type.value, section_title.split('>')[-1])
        
        # Revenir à la version simple qui fonctionnait auparavant
        task = OpenAIGenerationTask(path, prompt)

        task._finished_connection = task.signals.finished.connect(
            self.on_generation_finished
        )
        task._error_connection = task.signals.error.connect(self.on_generation_error)
        
        # Connecter le signal de progression si disponible
        if hasattr(task.signals, "progress"):
            task._progress_connection = task.signals.progress.connect(
                self.on_generation_progress
            )

        # Ajouter la tâche au pool de threads
        self.thread_pool.start(task)
        
        # Mettre à jour l'interface utilisateur
        self.generate_button.setText("Génération en cours…")
        self.generate_button.setEnabled(False)
        self.progress_bar.setValue(0)
        self.progress_bar.setVisible(True)
        self.status_label.setText(f"Génération en cours pour: {path}")
    
    def on_generation_finished(self, path, html):
        # Sauvegarder la version précédente si elle existe
        if path in self.generated_content:
            if path not in self.versions:
                self.versions[path] = {}
            
            self.versions[path][self.current_version] = self.generated_content[path]
            self.current_version += 1
        
        self.generated_content[path] = html
        if self.current_item_path == path:
            # Utiliser skip_title=True pour éviter la duplication des titres
            self.html_view.setHtml(render_html(html, self.default_css, skip_title=True))
            self.content_editor.setPlainText(html)
            self.generate_button.setText("Régénérer")
            self.generate_button.setEnabled(True)
            self.progress_bar.setVisible(False)
            self.status_label.setText(f"Génération terminée pour: {path}")
            
            # Mettre à jour le combo des versions
            self.update_version_combo(path)
        
        # Ajouter à l'historique
        self.add_to_history(path)
        
        # Sauvegarder automatiquement si activé
        if self.autosave_checkbox.isChecked():
            self.auto_save()

    def on_generation_error(self, msg):
        self.html_view.setHtml(f"<h2>Erreur</h2><p>{msg}</p>")
        self.generate_button.setText("Réessayer")
        self.generate_button.setEnabled(True)
        self.progress_bar.setVisible(False)
        self.status_label.setText(f"Erreur de génération: {msg}")
        
    def on_generation_progress(self, progress):
        self.progress_bar.setValue(progress)
        
    def generate_content_streaming(self):
        """Génère du contenu en utilisant le mode streaming de l'API OpenAI"""
        if not hasattr(self, "current_item_path") or not self.current_item_path:
            QMessageBox.warning(self, "Aucune section sélectionnée", "Veuillez sélectionner une section dans la table des matières.")
            return
        
        path = self.current_item_path
        
        # Réinitialiser le contenu en streaming
        self.streaming_content = ""
        self.current_streaming_path = path
        self.is_streaming = True
        # Préparer le prompt
        custom_prompt = self.custom_prompts.get(path)
        if custom_prompt:
            prompt = custom_prompt
        else:
            # Extraire le dernier élément du chemin pour le titre de la section
            section_title = path.split("<i class='nav-icon'></i>")[-1].strip() if "<i class='nav-icon'></i>" in path else path
            prompt = build_prompt(self.doc_type.value, section_title)
        
        # Créer une tâche de génération en streaming
        task = OpenAIStreamingTask(path, prompt)

        # Connecter les signaux
        task._finished_connection = task.signals.finished.connect(
            self.on_streaming_finished
        )
        task._error_connection = task.signals.error.connect(self.on_generation_error)
        task._chunk_connection = task.signals.chunk.connect(self.on_streaming_chunk)
        
        # Connecter le signal de progression si disponible
        if hasattr(task.signals, "progress"):
            task._progress_connection = task.signals.progress.connect(
                self.on_generation_progress
            )

        # Ajouter la tâche au pool de threads
        self.thread_pool.start(task)
        
        # Mettre à jour l'interface utilisateur
        self.generate_button.setText("Génération en cours…")
        self.generate_button.setEnabled(False)
        self.progress_bar.setValue(0)
        self.progress_bar.setVisible(True)
        self.status_label.setText(f"Génération en streaming pour: {path}")
    
    def on_streaming_chunk(self, chunk):
        """Gère la réception d'un nouveau morceau de texte généré en streaming"""
        # Mettre à jour le contenu en temps réel
        self.streaming_content += chunk
        
        # Mettre à jour l'affichage HTML si le chemin courant correspond
        if hasattr(self, "current_streaming_path") and self.current_item_path == self.current_streaming_path:
            # Préparer le contenu partiel pour l'affichage
            partial_html = self.streaming_content
            
            # Traiter les blocs de code incomplets
            import re
            
            # Fermer temporairement les balises de code ouvertes
            code_pattern = r'<pre><code[^>]*>((?:(?!</code></pre>).)*?)$'
            partial_html = re.sub(code_pattern, r'<pre><code\1</code></pre>', partial_html, flags=re.DOTALL)
            
            # Fermer temporairement les diagrammes Mermaid incomplets
            mermaid_pattern = r'<div class="mermaid">((?:(?!</div>).)*?)$'
            partial_html = re.sub(mermaid_pattern, r'<div class="mermaid">\1</div>', partial_html, flags=re.DOTALL)
            
            # Afficher le contenu partiel
            self.html_view.setHtml(render_html(partial_html, self.default_css, skip_title=True))
            self.content_editor.setPlainText(self.streaming_content)
            
            # Mettre à jour la progression
            self.progress_bar.setValue(50)  # Valeur arbitraire pour indiquer que ça avance
    
    def on_streaming_finished(self, path, content):
        """Gère la fin de la génération en streaming"""
        self.is_streaming = False
        self.generated_content[path] = content
        
        # Mettre à jour l'interface utilisateur
        if self.current_item_path == path:
            self.html_view.setHtml(render_html(content, self.default_css, skip_title=True))
            self.content_editor.setPlainText(content)
            self.generate_button.setText("Régénérer")
            self.generate_button.setEnabled(True)
            self.progress_bar.setVisible(False)
            self.status_label.setText(f"Génération terminée pour: {path}")
            
            # Mettre à jour le combo des versions
            self.update_version_combo(path)
        
        # Ajouter à l'historique
        self.add_to_history(path)
        
    def build_history_model(self):
        model = QStandardItemModel()
        model.setHorizontalHeaderLabels(["Historique des générations"])
        root = model.invisibleRootItem()
        
        for entry in self.history:
            timestamp = entry.get("timestamp", "")
            path = entry.get("path", "")
            display_text = f"{path} ({timestamp})"
            
            item = QStandardItem(display_text)
            item.setEditable(False)
            item.setData(path)
            root.appendRow(item)
        
        return model

    def get_full_path(self, index):
        parts = []
        model = index.model()
        while index.isValid():
            parts.insert(0, model.itemFromIndex(index).text())
            index = index.parent()
        return " > ".join(parts)

    def handle_tree_click(self, index):
        self.current_item_path = self.get_full_path(index)
        print(f"[DEBUG] current_item_path = {self.current_item_path}")
        if self.current_item_path:
            self.load_content(self.current_item_path)
        else:
            print("[WARNING] Aucun chemin sélectionné, chargement annulé.")
        self.add_to_history(self.current_item_path)
        
    def handle_favorite_click(self, index):
        path = index.data()
        if path:
            self.current_item_path = path
            self.load_content(path)
    
    def handle_history_click(self, index):
        path = index.model().itemFromIndex(index).data()
        if path:
            self.current_item_path = path
            self.load_content(path)
    
    def load_content(self, path):
        html = self.generated_content.get(path)
        if html:
            # Utiliser render_html pour appliquer les styles et icônes
            # Passer skip_title=True pour éviter d'ajouter un titre supplémentaire
            self.html_view.setHtml(render_html(html, self.default_css, skip_title=True))
            self.content_editor.setPlainText(html)
            self.generate_button.setText("Régénérer")
            self.update_version_combo(path)
        else:
            # Remplacer les symboles > par des icônes dans le titre de la page
            formatted_path = path
            if ">" in path:
                formatted_path = path.replace(" > ", " <i class='nav-icon'></i> ")
            
            # Utiliser la même police et style que le reste de l'application
            self.html_view.setHtml(render_html(
                f"<h1 class='page-title'>{formatted_path}</h1><p class='page-content'>Contenu non généré.</p>",
                self.default_css
            ))
            self.content_editor.clear()
            self.generate_button.setText("Générer")
            self.version_combo.clear()
        
        # Mettre à jour le statut avec des icônes au lieu des symboles >
        formatted_status_path = path
        if ">" in path:
            formatted_status_path = path.replace(" > ", " › ")
        self.status_label.setText(f"Section actuelle: {formatted_status_path}")

    def filter_toc(self, text):
        if not text:
            # Si le texte est vide, réinitialiser le modèle
            self.toc.setModel(self.build_tree_model())
            self.toc.expandAll()
            return
        
        # Créer un nouveau modèle pour les résultats filtrés
        filtered_model = QStandardItemModel()
        filtered_model.setHorizontalHeaderLabels(["Résultats de recherche"])
        root = filtered_model.invisibleRootItem()
        
        # Fonction récursive pour chercher dans le modèle original
        def search_in_model(parent_index, search_text):
            results = []
            model = self.toc.model()
            row_count = model.rowCount(parent_index)
            
            for row in range(row_count):
                index = model.index(row, 0, parent_index)
                text = model.data(index)
                
                # Vérifier si le texte correspond à la recherche
                if search_text.lower() in text.lower():
                    results.append((text, index))
                
                # Rechercher dans les enfants
                child_results = search_in_model(index, search_text)
                if child_results:
                    results.extend(child_results)
            
            return results
        
        # Effectuer la recherche
        results = search_in_model(QModelIndex(), text)
        
        # Ajouter les résultats au modèle filtré
        for text, _ in results:
            item = QStandardItem(text)
            item.setEditable(False)
            root.appendRow(item)
        
        # Appliquer le modèle filtré
        self.toc.setModel(filtered_model)

    def generate_content(self):
        if not hasattr(self, "current_item_path"):
            return

        path = self.current_item_path
        
        # Vérifier si un prompt personnalisé existe
        custom_prompt = self.custom_prompts.get(path)
        if custom_prompt:
            prompt = custom_prompt
        else:
            # Extraire le dernier élément du chemin pour le titre de la section
            section_title = path.split("<i class='nav-icon'></i>")[-1].strip() if "<i class='nav-icon'></i>" in path else path
            prompt = build_prompt(self.doc_type.value, section_title.split('>')[-1])
        
        # Revenir à la version simple qui fonctionnait auparavant
        task = OpenAIGenerationTask(path, prompt)

        task._finished_connection = task.signals.finished.connect(
            self.on_generation_finished
        )
        task._error_connection = task.signals.error.connect(self.on_generation_error)
        
        # Connecter le signal de progression si disponible
        if hasattr(task.signals, "progress"):
            task._progress_connection = task.signals.progress.connect(
                self.on_generation_progress
            )

        # Ajouter la tâche au pool de threads
        self.thread_pool.start(task)
        
        # Mettre à jour l'interface utilisateur
        self.generate_button.setText("Génération en cours…")
        self.generate_button.setEnabled(False)
        self.progress_bar.setValue(0)
        self.progress_bar.setVisible(True)
        self.status_label.setText(f"Génération en cours pour: {path}")

    def on_generation_finished(self, path, html):
        # Sauvegarder la version précédente si elle existe
        if path in self.generated_content:
            if path not in self.versions:
                self.versions[path] = {}
            
            self.versions[path][self.current_version] = self.generated_content[path]
            self.current_version += 1
        
        self.generated_content[path] = html
        if self.current_item_path == path:
            # Utiliser skip_title=True pour éviter la duplication des titres
            self.html_view.setHtml(render_html(html, self.default_css, skip_title=True))
            self.content_editor.setPlainText(html)
            self.generate_button.setText("Régénérer")
            self.generate_button.setEnabled(True)
            self.progress_bar.setVisible(False)
            self.status_label.setText(f"Génération terminée pour: {path}")
            
            # Mettre à jour le combo des versions
            self.update_version_combo(path)
        
        # Ajouter à l'historique
        self.add_to_history(path)
        
        # Sauvegarder automatiquement si activé
        if self.autosave_checkbox.isChecked():
            self.auto_save()

    def on_generation_error(self, msg):
        self.html_view.setHtml(f"<h2>Erreur</h2><p>{msg}</p>")
        self.generate_button.setText("Réessayer")
        self.generate_button.setEnabled(True)
        self.progress_bar.setVisible(False)
        self.status_label.setText(f"Erreur de génération: {msg}")
        
    def on_generation_progress(self, progress):
        self.progress_bar.setValue(progress)
        
    def generate_content_streaming(self):
        """Génère du contenu en utilisant le mode streaming de l'API OpenAI"""
        if not hasattr(self, "current_item_path") or not self.current_item_path:
            QMessageBox.warning(self, "Aucune section sélectionnée", "Veuillez sélectionner une section dans la table des matières.")
            return
        
        path = self.current_item_path
        
        # Réinitialiser le contenu en streaming
        self.streaming_content = ""
        self.current_streaming_path = path
        self.is_streaming = True
        
        # Préparer le prompt
        custom_prompt = self.custom_prompts.get(path)
        if custom_prompt:
            prompt = custom_prompt
        else:
            # Extraire le dernier élément du chemin pour le titre de la section
            section_title = path.split("<i class='nav-icon'></i>")[-1].strip() if "<i class='nav-icon'></i>" in path else path
            prompt = build_prompt(self.doc_type.value, section_title)
        
        # Créer une tâche de génération en streaming
        task = OpenAIStreamingTask(path, prompt)

        # Connecter les signaux
        task._finished_connection = task.signals.finished.connect(
            self.on_streaming_finished
        )
        task._error_connection = task.signals.error.connect(self.on_generation_error)
        task._chunk_connection = task.signals.chunk.connect(self.on_streaming_chunk)
        
        # Connecter le signal de progression si disponible
        if hasattr(task.signals, "progress"):
            task._progress_connection = task.signals.progress.connect(
                self.on_generation_progress
            )

        # Ajouter la tâche au pool de threads
        self.thread_pool.start(task)
        
        # Mettre à jour l'interface utilisateur
        self.generate_button.setText("Génération en cours…")
        self.generate_streaming_button.setText("Streaming en cours…")
        self.generate_button.setEnabled(False)
        self.generate_streaming_button.setEnabled(False)
        self.progress_bar.setValue(0)
        self.progress_bar.setVisible(True)
        self.status_label.setText(f"Génération en streaming pour: {path}")
    
    def on_streaming_finished(self, path, html):
        """Gère la fin de la génération en streaming"""
        # Sauvegarder la version précédente si elle existe
        if path in self.generated_content:
            if path not in self.versions:
                self.versions[path] = {}
            
            self.versions[path][self.current_version] = self.generated_content[path]
            self.current_version += 1
        
        # Enregistrer le contenu final
        self.generated_content[path] = html
        self.is_streaming = False
        
        if self.current_item_path == path:
            # Afficher le contenu final
            self.html_view.setHtml(render_html(html, self.default_css, skip_title=True))
            self.content_editor.setPlainText(html)
            self.generate_button.setText("Régénérer")
            self.generate_streaming_button.setText("Générer (Streaming)")
            self.generate_button.setEnabled(True)
            self.generate_streaming_button.setEnabled(True)
            self.progress_bar.setVisible(False)
            self.status_label.setText(f"Génération en streaming terminée pour: {path}")
            
            # Mettre à jour le combo des versions
            self.update_version_combo(path)
        
        # Ajouter à l'historique
        self.add_to_history(path)
        
        # Sauvegarder automatiquement si activé
        if self.autosave_checkbox.isChecked():
            self.auto_save()
    
    def on_streaming_chunk(self, chunk):
        """Gère la réception d'un nouveau morceau de texte généré en streaming"""
        # Mettre à jour le contenu en temps réel
        self.streaming_content += chunk
        
        # Mettre à jour l'affichage HTML si le chemin courant correspond
        if hasattr(self, "current_streaming_path") and self.current_item_path == self.current_streaming_path:
            # Préparer le contenu partiel pour l'affichage
            partial_html = self.streaming_content
            
            # Traiter les blocs de code incomplets
            import re
            
            # Fermer temporairement les balises de code ouvertes
            code_pattern = r'<pre><code[^>]*>((?:(?!</code></pre>).)*?)$'
            partial_html = re.sub(code_pattern, r'<pre><code\1</code></pre>', partial_html, flags=re.DOTALL)
            
            # Fermer temporairement les diagrammes Mermaid incomplets
            mermaid_pattern = r'<div class="mermaid">((?:(?!</div>).)*?)$'
            partial_html = re.sub(mermaid_pattern, r'<div class="mermaid">\1</div>', partial_html, flags=re.DOTALL)
            
            # Ajouter un script pour réinitialiser Mermaid à chaque mise à jour
            mermaid_reset_script = """
            <script>
            if (typeof mermaid !== 'undefined') {
                try {
                    mermaid.contentLoaded();
                } catch(e) { console.error(e); }
            }
            </script>
            """
            
            # Ajouter l'indicateur de frappe
            partial_html += "<span class='typing-indicator'>...</span>" + mermaid_reset_script
            
            # Rendre le HTML avec les balises temporairement fermées
            self.html_view.setHtml(render_html(partial_html, self.default_css, skip_title=True))
            
            # Mettre à jour l'éditeur de contenu avec le contenu réel (sans les fermetures temporaires)
            self.content_editor.setPlainText(self.streaming_content)
        
    def update_version_combo(self, path):
        self.version_combo.clear()
        if path in self.versions:
            for version in sorted(self.versions[path].keys()):
                self.version_combo.addItem(f"Version {version}", version)
            self.version_combo.addItem("Version actuelle", "current")
            self.version_combo.setCurrentIndex(self.version_combo.count() - 1)
    
    def load_version(self, index):
        if index < 0 or not hasattr(self, "current_item_path"):
            return
            
        version = self.version_combo.itemData(index)
        if version == "current":
            html = self.generated_content.get(self.current_item_path, "")
        else:
            html = self.versions.get(self.current_item_path, {}).get(version, "")
            
        if html:
            self.version_preview.setHtml(render_html(html))
    
    def apply_content_edits(self):
        if not hasattr(self, "current_item_path"):
            return
            
        html = self.content_editor.toPlainText()
        
        # Sauvegarder la version précédente
        if self.current_item_path in self.generated_content:
            if self.current_item_path not in self.versions:
                self.versions[self.current_item_path] = {}
                
            self.versions[self.current_item_path][self.current_version] = self.generated_content[self.current_item_path]
            self.current_version += 1
            
        # Mettre à jour le contenu
        self.generated_content[self.current_item_path] = html
        self.html_view.setHtml(render_html(html, self.default_css))
        
        # Mettre à jour les versions
        self.update_version_combo(self.current_item_path)
        
        # Mettre à jour le statut
        self.status_label.setText(f"Modifications appliquées pour: {self.current_item_path}")
        
        # Sauvegarder automatiquement si activé
        if self.autosave_checkbox.isChecked():
            self.auto_save()
    
    def edit_content(self):
        self.content_tabs.setCurrentIndex(1)  # Aller à l'onglet d'édition
    
    def add_to_favorites(self):
        if not hasattr(self, "current_item_path"):
            return
            
        if self.current_item_path not in self.favorites:
            self.favorites.append(self.current_item_path)
            self.favorites_list.setModel(self.build_favorites_model())
            self.save_favorites()
            self.status_label.setText(f"Ajouté aux favoris: {self.current_item_path}")
    
    def remove_from_favorites(self, path):
        if path in self.favorites:
            self.favorites.remove(path)
            self.favorites_list.setModel(self.build_favorites_model())
            self.save_favorites()
            self.status_label.setText(f"Retiré des favoris: {path}")
    
    def add_to_history(self, path):
        # Ajouter au début de l'historique
        timestamp = datetime.datetime.now().strftime("%d/%m/%Y %H:%M")
        entry = {"path": path, "timestamp": timestamp}
        
        # Supprimer les entrées existantes pour ce chemin
        self.history = [e for e in self.history if e.get("path") != path]
        
        # Ajouter la nouvelle entrée
        self.history.insert(0, entry)
        
        # Limiter la taille de l'historique
        if len(self.history) > 50:
            self.history = self.history[:50]
            
        # Mettre à jour le modèle
        self.history_list.setModel(self.build_history_model())
        
        # Sauvegarder l'historique
        self.save_history()
    
    def load_favorites(self):
        try:
            favorites_file = self.save_dir / f"favorites_{self.doc_type.name}.json"
            if favorites_file.exists():
                with open(favorites_file, "r", encoding="utf-8") as f:
                    return json.load(f)
        except Exception as e:
            print(f"Erreur lors du chargement des favoris: {e}")
        return []
    
    def save_favorites(self):
        try:
            favorites_file = self.save_dir / f"favorites_{self.doc_type.name}.json"
            with open(favorites_file, "w", encoding="utf-8") as f:
                json.dump(self.favorites, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Erreur lors de la sauvegarde des favoris: {e}")
    
    def load_history(self):
        try:
            history_file = self.save_dir / f"history_{self.doc_type.name}.json"
            if history_file.exists():
                with open(history_file, "r", encoding="utf-8") as f:
                    return json.load(f)
        except Exception as e:
            print(f"Erreur lors du chargement de l'historique: {e}")
        return []
    
    def save_history(self):
        try:
            history_file = self.save_dir / f"history_{self.doc_type.name}.json"
            with open(history_file, "w", encoding="utf-8") as f:
                json.dump(self.history, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Erreur lors de la sauvegarde de l'historique: {e}")
    
    def load_custom_prompts(self):
        try:
            prompts_file = self.save_dir / f"prompts_{self.doc_type.name}.json"
            if prompts_file.exists():
                with open(prompts_file, "r", encoding="utf-8") as f:
                    return json.load(f)
        except Exception as e:
            print(f"Erreur lors du chargement des prompts personnalisés: {e}")
        return {}
    
    def save_custom_prompts(self):
        try:
            prompts_file = self.save_dir / f"prompts_{self.doc_type.name}.json"
            with open(prompts_file, "w", encoding="utf-8") as f:
                json.dump(self.custom_prompts, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Erreur lors de la sauvegarde des prompts personnalisés: {e}")
    
    def customize_prompt(self):
        if not hasattr(self, "current_item_path"):
            return
            
        # Récupérer le prompt actuel
        current_prompt = self.custom_prompts.get(
            self.current_item_path,
            build_prompt(self.doc_type.value, self.current_item_path.split("<i class='nav-icon'></i>")[-1].strip() if "<i class='nav-icon'></i>" in self.current_item_path else self.current_item_path)
        )
        
        # Créer une nouvelle instance de PromptEditorDialog avec le prompt actuel
        dialog = PromptEditorDialog(self, current_prompt)
        
        # Exécuter le dialogue
        if dialog.exec():
            # Sauvegarder le prompt personnalisé
            self.custom_prompts[self.current_item_path] = dialog.get_prompt()
            self.save_custom_prompts()
            self.status_label.setText(f"Prompt personnalisé pour: {self.current_item_path}")
    
    def show_toc_context_menu(self, position):
        index = self.toc.indexAt(position)
        if not index.isValid():
            return
            
        path = self.get_full_path(index)
        
        menu = QMenu(self)
        
        generate_action = QAction("Générer", self)
        generate_action.triggered.connect(lambda: self.generate_for_path(path))
        menu.addAction(generate_action)
        
        if path in self.favorites:
            remove_favorite_action = QAction("Retirer des favoris", self)
            remove_favorite_action.triggered.connect(lambda: self.remove_from_favorites(path))
            menu.addAction(remove_favorite_action)
        else:
            add_favorite_action = QAction("Ajouter aux favoris", self)
            add_favorite_action.triggered.connect(lambda: self.add_to_favorites_path(path))
            menu.addAction(add_favorite_action)
        
        customize_prompt_action = QAction("Personnaliser le prompt", self)
        customize_prompt_action.triggered.connect(lambda: self.customize_prompt_for_path(path))
        menu.addAction(customize_prompt_action)
        
        menu.exec_(self.toc.viewport().mapToGlobal(position))
    
    def generate_for_path(self, path):
        self.current_item_path = path
        self.generate_content()
    
    def add_to_favorites_path(self, path):
        self.current_item_path = path
        self.add_to_favorites()
    
    def customize_prompt_for_path(self, path):
        self.current_item_path = path
        self.customize_prompt()
    
    def toggle_html_source(self):
        if not hasattr(self, "current_item_path"):
            return

        html = self.generated_content.get(self.current_item_path)
        if not html:
            return

        if self.html_view.page().title() == "Code source":
            self.html_view.setHtml(render_html(html, self.default_css))
            self.show_source_button.setText("Afficher le source")
        else:
            self.show_html_source()
            self.show_source_button.setText("Masquer le source")

    def show_html_source(self):
        if not hasattr(self, "current_item_path"):
            return

        html = self.generated_content.get(self.current_item_path)
        if not html:
            return

        # Extraire uniquement le dernier titre de l'arborescence
        title = (
            self.current_item_path.split("<i class='nav-icon'></i>")[-1].strip() if "<i class='nav-icon'></i>" in self.current_item_path else self.current_item_path
            if " > " in self.current_item_path
            else self.current_item_path
        )

        escaped_html = (
            html.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        )
        source_view = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Code source</title>
            <style>
                pre {{
                    background: #f5f5f5;
                    padding: 15px;
                    border-radius: 5px;
                    overflow-x: auto;
                    font-family: 'Courier New', monospace;
                    font-size: 14px;
                    line-height: 1.4;
                }}
                body {{
                    font-family: 'Segoe UI', sans-serif;
                    padding: 20px;
                }}
                h2 {{
                    color: #333;
                    border-bottom: 1px solid #ddd;
                    padding-bottom: 10px;
                }}
            </style>
        </head>
        <body>
            <h2>Code source : {title}</h2>
            <pre>{escaped_html}</pre>
        </body>
        </html>
        """
        self.html_view.setHtml(source_view)
    
    def export_to_pdf(self):
        if not hasattr(self, "current_item_path"):
            return

        html = self.generated_content.get(self.current_item_path)
        if not html:
            return

        if ">" in self.current_item_path:
            filename = (self.current_item_path.split("<i class='nav-icon'></i>")[-1] if "<i class='nav-icon'></i>" in self.current_item_path else self.current_item_path).replace(" ", "") + ".pdf"
        else:
            filename = self.current_item_path.replace(" ", "") + ".pdf"

        file_path, _ = QFileDialog.getSaveFileName(
            self, "Exporter en PDF", filename, "Fichiers PDF (*.pdf)"
        )

        if file_path:
            export_pdf(self.html_view, file_path)
            self.status_label.setText(f"Exporté en PDF: {file_path}")
    
    def export_to_word(self):
        if not hasattr(self, "current_item_path"):
            return
            
        html = self.generated_content.get(self.current_item_path)
        if not html:
            return
            
        if ">" in self.current_item_path:
            filename = (self.current_item_path.split("<i class='nav-icon'></i>")[-1] if "<i class='nav-icon'></i>" in self.current_item_path else self.current_item_path).replace(" ", "") + ".docx"
        else:
            filename = self.current_item_path.replace(" ", "") + ".docx"
            
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Exporter en Word", filename, "Fichiers Word (*.docx)"
        )
        
        if file_path:
            try:
                # Utiliser python-docx pour créer un document Word
                from docx import Document
                from docx.shared import Pt
                
                doc = Document()
                doc.add_heading(self.current_item_path, 0)
                
                # Ajouter le contenu (simplifié, sans le HTML)
                doc.add_paragraph(html)
                
                doc.save(file_path)
                self.status_label.setText(f"Exporté en Word: {file_path}")
            except Exception as e:
                QMessageBox.warning(self, "Erreur d'exportation", f"Erreur lors de l'exportation en Word: {e}")
    
    def export_to_markdown(self):
        if not hasattr(self, "current_item_path"):
            return
            
        html = self.generated_content.get(self.current_item_path)
        if not html:
            return
            
        if ">" in self.current_item_path:
            filename = (self.current_item_path.split("<i class='nav-icon'></i>")[-1] if "<i class='nav-icon'></i>" in self.current_item_path else self.current_item_path).replace(" ", "") + ".md"
        else:
            filename = self.current_item_path.replace(" ", "") + ".md"
            
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Exporter en Markdown", filename, "Fichiers Markdown (*.md)"
        )
        
        if file_path:
            try:
                # Conversion HTML vers Markdown simplifiée
                try:
                    # Essayer d'utiliser html2text si disponible
                    import html2text
                    converter = html2text.HTML2Text()
                    converter.ignore_links = False
                    converter.ignore_images = False
                    markdown = converter.handle(html)
                except ImportError:
                    # Si html2text n'est pas disponible, utiliser une conversion basique
                    QMessageBox.information(
                        self, 
                        "Module manquant", 
                        "Le module html2text n'est pas installé. Une conversion simplifiée sera utilisée.\n"
                        "Pour une meilleure conversion, installez html2text avec: pip install html2text"
                    )
                    # Conversion basique
                    markdown = html
                    # Supprimer les balises HTML les plus courantes
                    import re
                    markdown = re.sub(r'<[^>]*>', '', markdown)
                    # Ajouter un titre
                    markdown = f"# {self.current_item_path}\n\n{markdown}"
                
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(markdown)
                    
                self.status_label.setText(f"Exporté en Markdown: {file_path}")
            except Exception as e:
                QMessageBox.warning(self, "Erreur d'exportation", f"Erreur lors de l'exportation en Markdown: {e}")
    
    def save_document(self):
        if not self.generated_content:
            QMessageBox.information(self, "Information", "Aucun contenu à sauvegarder.")
            return
            
        # Créer un dictionnaire avec les données à sauvegarder
        data = {
            "doc_type": self.doc_type.name,
            "content": self.generated_content,
            "versions": self.versions,
            "timestamp": datetime.datetime.now().strftime("%d/%m/%Y %H:%M")
        }
        
        # Demander le nom du fichier
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Sauvegarder le document", 
            str(self.save_dir / f"{self.doc_type.name}.json"),
            "Fichiers JSON (*.json)"
        )
        
        if file_path:
            try:
                with open(file_path, "w", encoding="utf-8") as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                    
                self.last_save_label.setText(f"Dernière sauvegarde: {data['timestamp']}")
                self.status_label.setText(f"Document sauvegardé: {file_path}")
            except Exception as e:
                QMessageBox.warning(self, "Erreur de sauvegarde", f"Erreur lors de la sauvegarde: {e}")
    
    def load_document(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Charger un document", 
            str(self.save_dir),
            "Fichiers JSON (*.json)"
        )
        
        if file_path:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    
                # Vérifier que le type de document correspond
                if data.get("doc_type") != self.doc_type.name:
                    QMessageBox.warning(
                        self, 
                        "Type de document incorrect", 
                        f"Ce document est de type {data.get('doc_type')} et non {self.doc_type.name}."
                    )
                    return
                    
                # Charger les données
                self.generated_content = data.get("content", {})
                self.versions = data.get("versions", {})
                
                # Mettre à jour l'interface
                if hasattr(self, "current_item_path"):
                    self.load_content(self.current_item_path)
                    
                timestamp = data.get("timestamp", "Inconnu")
                self.last_save_label.setText(f"Dernière sauvegarde: {timestamp}")
                self.status_label.setText(f"Document chargé: {file_path}")
            except Exception as e:
                QMessageBox.warning(self, "Erreur de chargement", f"Erreur lors du chargement: {e}")
    
    def auto_save(self):
        if not self.generated_content or not self.autosave_checkbox.isChecked():
            return
            
        # Créer un dictionnaire avec les données à sauvegarder
        data = {
            "doc_type": self.doc_type.name,
            "content": self.generated_content,
            "versions": self.versions,
            "timestamp": datetime.datetime.now().strftime("%d/%m/%Y %H:%M")
        }
        
        # Sauvegarder dans un fichier auto-save
        file_path = self.save_dir / f"autosave_{self.doc_type.name}.json"
        
        try:
            with open(file_path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
                
            self.last_save_label.setText(f"Dernière sauvegarde: {data['timestamp']}")
            self.status_label.setText("Sauvegarde automatique effectuée")
        except Exception as e:
            print(f"Erreur lors de la sauvegarde automatique: {e}")
    
    def toggle_autosave(self, state):
        if state == Qt.Checked:
            self.autosave_timer.start(60000)
            self.status_label.setText("Sauvegarde automatique activée")
        else:
            self.autosave_timer.stop()
            self.status_label.setText("Sauvegarde automatique désactivée")
            
    def publish_to_git(self):
        """Publier le contenu généré sur un dépôt Git"""
        if not self.generated_content:
            QMessageBox.information(self, "Information", "Aucun contenu à publier.")
            return
        
        # Demander les informations Git en utilisant le composant lazy loaded
        dialog = self.git_credentials_dialog.get()
        if not dialog.exec():
            return
        
        credentials = dialog.get_credentials()
        
        # Vérifier que les champs obligatoires sont remplis
        if not credentials["repo"]:
            QMessageBox.warning(self, "Erreur", "L'URL du dépôt Git est obligatoire.")
            return
        
        # Créer un dossier temporaire pour la publication
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                self.status_label.setText("Préparation des fichiers pour publication...")
                
                # Exporter tous les contenus générés en HTML et Markdown
                for path, content in self.generated_content.items():
                    # Créer un nom de fichier valide
                    filename = path.replace(" > ", "_").replace(" ", "_")
                    
                    # Exporter en HTML
                    html_path = os.path.join(temp_dir, f"{filename}.html")
                    with open(html_path, "w", encoding="utf-8") as f:
                        f.write(render_html(content))
                    
                    # Exporter en Markdown (si possible)
                    try:
                        import html2text
                        converter = html2text.HTML2Text()
                        converter.ignore_links = False
                        converter.ignore_images = False
                        markdown = converter.handle(content)
                        
                        md_path = os.path.join(temp_dir, f"{filename}.md")
                        with open(md_path, "w", encoding="utf-8") as f:
                            f.write(markdown)
                    except ImportError:
                        pass
                
                # Créer un fichier README.md
                readme_path = os.path.join(temp_dir, "README.md")
                with open(readme_path, "w", encoding="utf-8") as f:
                    f.write(f"# Documentation {self.doc_type.value}\n\n")
                    f.write("Ce dépôt contient la documentation générée par AssistantPM.\n\n")
                    f.write("## Table des matières\n\n")
                    
                    for path in self.generated_content.keys():
                        filename = path.replace(" > ", "_").replace(" ", "_")
                        f.write(f"- [{path}]({filename}.md)\n")
                
                # Initialiser le dépôt Git
                self.status_label.setText("Initialisation du dépôt Git...")
                
                # Vérifier si git est installé
                try:
                    subprocess.run(["git", "--version"], check=True, capture_output=True)
                except (subprocess.CalledProcessError, FileNotFoundError):
                    QMessageBox.critical(self, "Erreur", "Git n'est pas installé ou n'est pas accessible.")
                    return
                
                # Initialiser le dépôt
                try:
                    # Initialiser le dépôt
                    subprocess.run(["git", "init"], cwd=temp_dir, check=True, capture_output=True)
                    
                    # Configurer l'utilisateur Git si nécessaire
                    if credentials["username"]:
                        subprocess.run(["git", "config", "user.name", credentials["username"]], 
                                      cwd=temp_dir, check=True, capture_output=True)
                        subprocess.run(["git", "config", "user.email", f"{credentials['username']}@users.noreply.github.com"], 
                                      cwd=temp_dir, check=True, capture_output=True)
                    
                    # Ajouter tous les fichiers
                    subprocess.run(["git", "add", "."], cwd=temp_dir, check=True, capture_output=True)
                    
                    # Créer un commit
                    commit_message = f"Publication de la documentation {self.doc_type.value}"
                    subprocess.run(["git", "commit", "-m", commit_message], cwd=temp_dir, check=True, capture_output=True)
                    
                    # Ajouter le dépôt distant
                    remote_url = credentials["repo"]
                    if credentials["username"] and credentials["password"]:
                        # Insérer les identifiants dans l'URL
                        if remote_url.startswith("https://"):
                            auth_part = f"{credentials['username']}:{credentials['password']}@"
                            remote_url = remote_url.replace("https://", f"https://{auth_part}")
                    
                    subprocess.run(["git", "remote", "add", "origin", remote_url], 
                                  cwd=temp_dir, check=True, capture_output=True)
                    
                    # Pousser vers le dépôt distant
                    self.status_label.setText("Publication en cours...")
                    branch = credentials["branch"] or "main"
                    
                    # Utiliser -u pour définir la branche amont
                    result = subprocess.run(
                        ["git", "push", "-u", "origin", branch], 
                        cwd=temp_dir, 
                        capture_output=True,
                        text=True
                    )
                    
                    if result.returncode != 0:
                        # Si la branche n'existe pas, essayer de la créer
                        if "fatal: la branche" in result.stderr or "fatal: The current branch" in result.stderr:
                            subprocess.run(["git", "checkout", "-b", branch], cwd=temp_dir, check=True)
                            subprocess.run(["git", "push", "-u", "origin", branch], cwd=temp_dir, check=True)
                        else:
                            raise subprocess.CalledProcessError(result.returncode, result.args, result.stdout, result.stderr)
                    
                    QMessageBox.information(
                        self, 
                        "Publication réussie", 
                        f"La documentation a été publiée avec succès sur {credentials['repo']}."
                    )
                    self.status_label.setText(f"Publication réussie sur {credentials['repo']}")
                    
                except subprocess.CalledProcessError as e:
                    error_message = f"Erreur lors de la publication: {e.stderr if hasattr(e, 'stderr') else str(e)}"
                    QMessageBox.critical(self, "Erreur de publication", error_message)
                    self.status_label.setText("Échec de la publication")
                    
        except Exception as e:
            QMessageBox.critical(self, "Erreur", f"Une erreur est survenue: {str(e)}")
            self.status_label.setText("Échec de la publication")
    
    def generate_all_content(self):
        # Récupérer tous les chemins possibles
        paths = []
        
        def collect_paths(parent_index, current_path=""):
            model = self.toc.model()
            row_count = model.rowCount(parent_index)
            
            for row in range(row_count):
                index = model.index(row, 0, parent_index)
                text = model.data(index)
                
                # Utilisation d'une icône au lieu du symbole '>' pour la navigation
                path = text if not current_path else f"{current_path} <i class='nav-icon'></i> {text}"
                paths.append(path)
                
                # Récursivement collecter les chemins des enfants
                collect_paths(index, path)
        
        collect_paths(QModelIndex())
        
        if not paths:
            return
            
        # Demander confirmation
        reply = QMessageBox.question(
            self, 
            "Générer tout le contenu", 
            f"Voulez-vous générer le contenu pour {len(paths)} sections ?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        
        if reply == QMessageBox.No:
            return
        
        # Préparer la barre de progression
        total_sections = len(paths)
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        self.progress_bar.setVisible(True)
        self.status_label.setText("Génération de la documentation complète...")
        
        # Générer le contenu pour chaque chemin
        for i, path in enumerate(paths):
            # Mettre à jour la barre de progression
            progress_percent = int((i / total_sections) * 100)
            self.progress_bar.setValue(progress_percent)
            self.status_label.setText(f"Génération en cours: {path} ({i+1}/{total_sections})")
            
            # Générer le contenu
            self.current_item_path = path
            self.generate_content()
            QApplication.processEvents()  # Permettre à l'interface de se mettre à jour
        
        # Finaliser la barre de progression
        self.progress_bar.setValue(100)
        self.status_label.setText(f"Génération complète: {total_sections} sections générées")


if __name__ == "__main__":
    import sys
    from PySide6.QtWidgets import QApplication
    from project.documents.DocType import DocType

    app = QApplication(sys.argv)

    # Créer et afficher le widget avec un type de document valide
    widget = DocumentationWidget(DocType.CAHIER_DES_CHARGES_FONCTIONNEL)
    widget.show()

    sys.exit(app.exec())
